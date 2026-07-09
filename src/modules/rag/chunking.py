import os
import re
import logging
import sys

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "..")))

from docx import Document as DocxDocument
from langchain.schema import Document
from langchain_core.documents import Document as LCDocument

from src.clients.embedding import embeddings_qa
from src.clients.llm import LLMClient
from src.configs import env_config
from src.state import extract_schema_key_word_chunking
from prompts.prompt import KEYWORD_EXTRACT_PROMPT
from vectorstores import VectorStoreManager

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)



print("ok ban oi")

# ── Constants ─────────────────────────────────────────────────────────────────
DOC_DIR = r"D:\VKU\Nam_4\ky_I\computer_vision\EDUAGENT\src\modules\rag\documents\doc"
COLLECTION_NAME = "doc"
DISEASE_PATTERN = r"(#[^#]*?(?=\n#|$))"

# ── LLM ───────────────────────────────────────────────────────────────────────
llm_keyword_extractor = LLMClient(
    model=env_config.model,
    api_provider=env_config.api_provider,
)._llm.with_structured_output(extract_schema_key_word_chunking)


# ── Step 1: Đọc tất cả file .docx trong thư mục ───────────────────────────────
def load_all_docx(doc_dir: str) -> list[str]:
    """Đọc tất cả file .docx, trả về list raw text."""
    all_texts = []
    docx_files = [f for f in os.listdir(doc_dir) if f.endswith(".docx")]

    if not docx_files:
        logger.warning(f"Không tìm thấy file .docx nào trong: {doc_dir}")
        return all_texts

    for filename in docx_files:
        path = os.path.join(doc_dir, filename)
        try:
            doc = DocxDocument(path)
            full_text = "\n".join(para.text for para in doc.paragraphs)
            all_texts.append(full_text)
            logger.info(f"✅ Đã đọc: {filename}")
        except Exception as e:
            logger.error(f"❌ Lỗi đọc {filename}: {e}")

    return all_texts


# ── Step 2: Tách từng bệnh theo pattern # ─────────────────────────────────────
def extract_disease_sections(full_text: str) -> list[str]:
    """Tách nội dung theo từng bệnh (bắt đầu bằng #)."""
    matches = re.findall(DISEASE_PATTERN, "\n" + full_text, re.DOTALL)
    return [m.strip() for m in matches if m.strip()]


# ── Step 3: Tách từng bệnh thành các chunks theo chủ đề ──────────────────────
def split_disease_into_chunks(disease_text: str) -> list[str]:
    """Tách 1 bệnh thành nhiều chunks theo đoạn văn."""
    lines = disease_text.strip().split("\n")
    disease_name = lines[0].replace("#", "").strip()
    body = "\n".join(lines[1:]).strip()

    chunks = []
    for para in body.split("\n\n"):
        para = para.strip()
        if para:
            chunks.append(f"[{disease_name}] - {para}")

    return chunks


# ── Step 4: Convert chunk → Document với metadata + keywords ─────────────────
def convert_chunk_to_document(chunk_text: str) -> tuple[str, dict]:
    """Parse chunk thành page_content và metadata có keywords từ LLM."""
    lines = [l.strip() for l in chunk_text.split("\n") if l.strip()]
    first_line = lines[0]

    bracket_match = re.search(r'\[(.+?)\]', first_line)
    bracket_content = bracket_match.group(1) if bracket_match else ""

    sci_match = re.search(r'\((.+?)\)', bracket_content)
    scientific_name = sci_match.group(1).strip() if sci_match else ""
    disease = re.sub(r'\(.*?\)', '', bracket_content).strip()

    topic_match = re.search(r'\]\s*-\s*(.+)', first_line)
    topic = topic_match.group(1).strip() if topic_match else ""

    # LLM extract keywords
    prompt = KEYWORD_EXTRACT_PROMPT.format(
        disease=disease,
        scientific_name=scientific_name,
        topic=topic,
        content=chunk_text,
    )
    response = llm_keyword_extractor.invoke(prompt)
    keywords = response.keywords if response else []

    page_content = (
        f"[{bracket_content}]\n"
        f"[{topic}]\n"
        + "\n".join(lines[1:])
    )

    metadata = {
        "disease": disease,
        "scientific_name": scientific_name,
        "topic": topic,
        "keywords": keywords,
    }

    return page_content, metadata


# ── Main pipeline ─────────────────────────────────────────────────────────────
def build_documents() -> list[Document]:
    """Pipeline đầy đủ: đọc docx → chunk → Document."""
    all_texts = load_all_docx(DOC_DIR)

    all_chunks = []
    for text in all_texts:
        sections = extract_disease_sections(text)
        for section in sections:
            all_chunks.extend(split_disease_into_chunks(section))

    logger.info(f"Tổng số chunks: {len(all_chunks)}")

    documents = []
    for i, chunk in enumerate(all_chunks):
        try:
            page_content, metadata = convert_chunk_to_document(chunk)
            documents.append(Document(page_content=page_content, metadata=metadata))
        except Exception as e:
            logger.error(f"❌ Lỗi chunk {i}: {e}")

    logger.info(f"Tổng số documents: {len(documents)}")
    return documents


def index_to_vector_store(documents: list[Document]) -> None:
    """Đưa documents vào Qdrant vector store."""
    vector_manager = VectorStoreManager(
        url=env_config.qdrant_url,
        api_key=env_config.qdrant_api_key,
    )
    vector_manager.create_vector_store(
        documents=documents,
        embeddings=embeddings_qa,
        collection_name=COLLECTION_NAME,
    )
    logger.info(f"✅ Đã index {len(documents)} documents vào collection '{COLLECTION_NAME}'")


# # ── Entry point ───────────────────────────────────────────────────────────────
# if __name__ == "__main__":
#     documents = build_documents()
#     index_to_vector_store(documents)
# # Chạy thử ngay - thêm print debug vào đầu file
# print(">>> Script started")

# Và đảm bảo cuối file có
if __name__ == "__main__":
    print(">>> Running main...")
    documents = build_documents()
    index_to_vector_store(documents)