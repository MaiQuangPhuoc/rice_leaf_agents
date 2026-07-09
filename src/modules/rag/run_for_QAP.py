from loaders import DocumentLoader
import logging
import sys , os 
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..","..")))

# from processors import DocumentProcessor
import re
from langchain_core.documents import Document
from vectorstores2 import VectorStoreManager

# from retrievers import VectorStoreRetriever
# from langchain.text_splitter import RecursiveCharacterTextSplitter


from src.clients.embedding import embeddings_qa
from src.configs import env_config

# from src.clients.llm import LLMClient
# from src.configs import env_config
from src.state import State, extract_schema_key_word_chunking
# from src.clients.databases import qdrant_qa

# from prompts.prompt import KEYWORD_EXTRACT_PROMPT

# from langchain.embeddings import HuggingFaceBgeEmbeddings


from docx import Document
import re

# llm_keyword_extractor = LLMClient(
#     model=env_config.model,
#     api_provider=env_config.api_provider
# )._llm.with_structured_output(extract_schema_key_word_chunking)


print("okoko")

file_path = r"D:\VKU\Nam_4\ky_I\computer_vision\EDUAGENT\src\modules\rag\documents\doc\khoe_manh.docx"
KEY_FILE = r"D:\VKU\Nam_4\ky_I\computer_vision\EDUAGENT\src\modules\key\khoe_manh.txt"
# Đọc hết nội dung file .docx thành 1 chuỗi text lớn
doc = Document(file_path)
full_text = "\n".join(para.text for para in doc.paragraphs)

# Tách theo pattern: bắt đầu bằng #Bệnh ... cho đến trước #Bệnh tiếp theo hoặc cuối file
pattern = r"(#[^#]*?(?=\n#|$))"
matches = re.findall(pattern, "\n" + full_text, re.DOTALL)

# Làm sạch từng phần (loại bỏ khoảng trắng thừa đầu/cuối)
docs = [match.strip() for match in matches if match.strip()]

def split_doc_into_chunks(doc: str):
    lines = doc.strip().split("\n")

    # lấy tên bệnh
    disease = lines[0].replace("#", "").strip()

    # phần nội dung còn lại
    body = "\n".join(lines[1:]).strip()

    # tách đoạn theo dòng trống
    raw_paragraphs = body.split("\n\n")

    chunks = []
    for idx, para in enumerate(raw_paragraphs, start=1):
        para = para.strip()
        if not para:
            continue
        
        chunk_title = f"[{disease}]"
        chunk = f"{chunk_title} - {para}"
        chunks.append(chunk)

    return chunks

all_chunks = []
for doc in docs:
    all_chunks.extend(split_doc_into_chunks(doc))

# for i, chunk in enumerate(all_chunks): 
#     print(f"\n--- Chunk {i+1} ---")
#     print(chunk, "\n----------------------------------------------------\n")  # in ra 500 ký tự đầu của chunk



def load_all_keyword_lists(file_path: str) -> list[list[str]]:
    """Đọc file txt, trả về list of list - mỗi keywords_n là 1 phần tử."""
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()

    matches = re.findall(r'keywords_\d+\s*=\s*\[([^\]]+)\]', content)

    all_keyword_lists = []
    for match in matches:
        items = re.findall(r"['\"](.+?)['\"]", match)
        all_keyword_lists.append(items)

    return all_keyword_lists


# Load 1 lần bên ngoài
keyword_lists = load_all_keyword_lists(KEY_FILE)


def convert_chunk_to_document(chunk_text: str, chunk_index: int):
    lines = [l.strip() for l in chunk_text.split("\n") if l.strip()]
    first_line = lines[0]

    bracket_match = re.search(r'\[(.+?)\]', first_line)
    bracket_content = bracket_match.group(1) if bracket_match else ""

    sci_match = re.search(r'\((.+?)\)', bracket_content)
    scientific_name = sci_match.group(1).strip() if sci_match else ""
    disease = re.sub(r'\(.*?\)', '', bracket_content).strip()

    topic_match = re.search(r'\]\s*-\s*(.+)', first_line)
    topic = topic_match.group(1).strip() if topic_match else ""

    content_lines = lines[1:]
    page_content = (
        f"{bracket_content} - {topic} {chr(10).join(content_lines)}\n")

    # Lấy keywords theo đúng thứ tự chunk
    keywords = keyword_lists[chunk_index] if chunk_index < len(keyword_lists) else []

    metadata = {
        "disease": disease,
        "scientific_name": scientific_name,
        "topic": topic,
        "keywords": keywords,
    }

    return page_content, metadata
 


 


# def convert_chunk_to_document(chunk_text: str):
#     lines = [l.strip() for l in chunk_text.split("\n") if l.strip()]

#     first_line = lines[0]

#     bracket_match = re.search(r'\[(.+?)\]', first_line)
#     bracket_content = bracket_match.group(1) if bracket_match else ""

#     sci_match = re.search(r'\((.+?)\)', bracket_content)
#     scientific_name = sci_match.group(1).strip() if sci_match else ""
#     disease = re.sub(r'\(.*?\)', '', bracket_content).strip()

#     topic_match = re.search(r'\]\s*-\s*(.+)', first_line)
#     topic = topic_match.group(1).strip() if topic_match else ""

#     # # Extract keywords
#     # prompt = KEYWORD_EXTRACT_PROMPT.format(
#     #     disease=disease,
#     #     scientific_name=scientific_name,
#     #     topic=topic,
#     #     content=chunk_text
#     # )

#     # response = llm_keyword_extractor.invoke(prompt)
    
#     # print(f'prompt: {prompt}')
#     # print(f'response: {response}\n\n ----------------------------------------------\n\n')
#     # keywords = response.keywords if response else []

#     # Build page_content
#     content_lines = lines[1:]
#     page_content = (
#         f"[{bracket_content}]\n"
#         f"[{topic}]\n"
#         + "\n".join(content_lines)
#     )

#     keywords = ['keyword1', 'keyword2']  # <-- tạm thời để trống, chờ phần trích xuất keyword hoạt động ổn định
#     metadata = {
#         "disease": disease,
#         "scientific_name": scientific_name,
#         "topic": topic,
#         "keywords": keywords,
#     }

#     return page_content, metadata


from langchain.schema import Document

documents = []

# for chunk in all_chunks:
#     page_content, metadata = convert_chunk_to_document(chunk)
#     documents.append(
#         Document(
#             page_content=page_content,
#             metadata=metadata
#         )
#     )

for i, chunk in enumerate(all_chunks):
    page_content, metadata = convert_chunk_to_document(chunk, chunk_index=i)
    documents.append(
        Document(
            page_content=page_content,
            metadata=metadata
        )
    )

for i, doc in enumerate(documents):
    print(f"\n--- Document {i+1} ---")
    print("Metadata:", doc.metadata)
    print("Page Content:", doc.page_content)


vector_manager = VectorStoreManager(
    url = env_config.qdrant_url,
    api_key=env_config.qdrant_api_key
)


vector_store = vector_manager.create_vector_store(
    documents=documents,
    embeddings=embeddings_qa,
    collection_name="documents"    
)

print("Vector store created successfully with the provided documents and embeddings.")